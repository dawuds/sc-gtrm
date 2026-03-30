#!/usr/bin/env python3
"""Generate SC-GTRM assessment report template (.docx)

Reads controls/library.json, controls/domains.json, and AWP data to produce
a professional Word document aligned to SC-GL/6-2023 regulatory submission
requirements for independent assessment reports.
"""

import json
import glob
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SCRIPT_DIR = Path(__file__).parent
PROJECT_DIR = SCRIPT_DIR.parent
OUTPUT = SCRIPT_DIR / "SC-GTRM-Report-Template.docx"

# Colours
NAVY = RGBColor(0x1F, 0x38, 0x64)
DARK_BLUE = RGBColor(0x2B, 0x4C, 0x7E)
MEDIUM_BLUE = RGBColor(0x44, 0x72, 0xC4)
DARK_GREY = RGBColor(0x33, 0x33, 0x33)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
NAVY_HEX = "1F3864"
LIGHT_BLUE_HEX = "D6E4F0"
WHITE_HEX = "FFFFFF"

# Domain ordering (matches SC-GL/6-2023 structure)
DOMAIN_ORDER = [
    "governance-oversight",
    "technology-risk-assessment",
    "cybersecurity",
    "cloud-services",
    "third-party-risk",
    "data-management",
    "incident-management",
    "business-continuity",
    "technology-audit",
    "emerging-technology",
]


def load_data():
    """Load controls, domains, and AWP metadata."""
    with open(PROJECT_DIR / "controls" / "library.json") as f:
        library = json.load(f)
    with open(PROJECT_DIR / "controls" / "domains.json") as f:
        domains = json.load(f)

    # Build AWP lookup: slug -> {ref, riskTier, domain}
    awp_lookup = {}
    for awp_file in sorted(glob.glob(str(SCRIPT_DIR / "awp-data" / "*.json"))):
        with open(awp_file) as f:
            awp_data = json.load(f)
        for section in awp_data.get("sections", []):
            awp_lookup[section["controlSlug"]] = {
                "ref": section["ref"],
                "riskTier": section["riskTier"],
                "awpDomain": section["domain"],
            }

    # Group controls by domain
    controls_by_domain = {}
    for ctrl in library["controls"]:
        domain_slug = ctrl["domain"]
        if domain_slug not in controls_by_domain:
            controls_by_domain[domain_slug] = []
        controls_by_domain[domain_slug].append(ctrl)

    return library, domains, awp_lookup, controls_by_domain


def configure_styles(doc):
    """Set up document styles for professional formatting."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = DARK_GREY
    pf = style.paragraph_format
    pf.space_after = Pt(4)
    pf.space_before = Pt(2)

    # Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = NAVY
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = DARK_BLUE
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)

    # Heading 3
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(11)
    h3.font.bold = True
    h3.font.color.rgb = MEDIUM_BLUE
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)

    return doc


def set_cell_shading(cell, colour_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{colour_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, font_size=Pt(9), colour=DARK_GREY,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(str(text))
    run.font.name = "Calibri"
    run.font.size = font_size
    run.font.bold = bold
    run.font.color.rgb = colour


def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a table with navy header row and alternating row shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    hdr_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_shading(cell, NAVY_HEX)
        set_cell_text(cell, header, bold=True, font_size=Pt(9), colour=WHITE,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = LIGHT_BLUE_HEX if r_idx % 2 == 0 else WHITE_HEX
        for c_idx, value in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_shading(cell, bg)
            set_cell_text(cell, value, font_size=Pt(9))

    # Apply column widths if specified
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = width

    return table


def add_page_break(doc):
    """Insert a page break."""
    doc.add_page_break()


def add_header_footer(doc):
    """Add CONFIDENTIAL header and page numbers footer to all sections."""
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run("CONFIDENTIAL")
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

        # Footer with page numbers
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run("SC-GTRM Independent Assessment Report  |  Page ")
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.color.rgb = DARK_GREY
        # Insert PAGE field
        fld_char_begin = parse_xml(
            f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'
        )
        instr_text = parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
        )
        fld_char_end = parse_xml(
            f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'
        )
        run2 = fp.add_run()
        run2._r.append(fld_char_begin)
        run2._r.append(instr_text)
        run2._r.append(fld_char_end)


def set_margins(doc):
    """Set 2.5cm margins on all sections."""
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def add_cover_page(doc):
    """Create the cover page."""
    # Spacing before title
    for _ in range(6):
        doc.add_paragraph("")

    # Classification
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONFIDENTIAL (SULIT)")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    doc.add_paragraph("")

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Independent Assessment Report")
    run.font.name = "Calibri"
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = NAVY

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SC-GL/6-2023 Guidelines on Technology Risk Management")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_BLUE

    doc.add_paragraph("")

    # Horizontal line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 60)
    run.font.color.rgb = MEDIUM_BLUE

    doc.add_paragraph("")

    # Cover details
    cover_fields = [
        ("Entity:", "[Entity Name]"),
        ("Prepared by:", "[Assessor Firm]"),
        ("Assessment Date:", "[DD/MM/YYYY]"),
        ("Engagement Reference:", "[Engagement Ref]"),
    ]

    table = doc.add_table(rows=len(cover_fields), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(cover_fields):
        # Remove borders by not using Table Grid style
        left_cell = table.rows[i].cells[0]
        right_cell = table.rows[i].cells[1]
        set_cell_text(left_cell, label, bold=True, font_size=Pt(11), colour=NAVY,
                      alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_text(right_cell, value, font_size=Pt(11), colour=DARK_GREY)
        left_cell.width = Cm(6)
        right_cell.width = Cm(9)

    # Remove table borders for cover page
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}/>'
    )
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tbl_pr.append(borders)


def add_document_control(doc):
    """Add document control section."""
    doc.add_heading("Document Control", level=1)

    # Version history
    doc.add_heading("Version History", level=2)
    headers = ["Version", "Date", "Author", "Changes"]
    rows = [
        ["0.1", "[DD/MM/YYYY]", "[Author Name]", "Initial draft"],
        ["0.2", "[DD/MM/YYYY]", "[Author Name]", "[Description of changes]"],
        ["1.0", "[DD/MM/YYYY]", "[Author Name]", "Final report issued"],
    ]
    add_styled_table(doc, headers, rows)

    doc.add_paragraph("")

    # Distribution list
    doc.add_heading("Distribution List", level=2)
    headers = ["Name", "Role", "Organisation"]
    rows = [
        ["[Name]", "[Role]", "[Organisation]"],
        ["[Name]", "[Role]", "[Organisation]"],
        ["[Name]", "[Role]", "[Organisation]"],
    ]
    add_styled_table(doc, headers, rows)


def add_toc(doc):
    """Add a Table of Contents field."""
    doc.add_heading("Table of Contents", level=1)

    p = doc.add_paragraph()
    run = p.add_run()
    fld_char_begin = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'
    )
    instr_text = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve">'
        f' TOC \\o "1-3" \\h \\z \\u </w:instrText>'
    )
    fld_char_separate = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'
    )
    fld_char_end = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'
    )
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)

    # Placeholder text
    run2 = p.add_run(
        "[Right-click and select 'Update Field' to generate Table of Contents]"
    )
    run2.font.name = "Calibri"
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run2.font.italic = True

    run3 = p.add_run()
    run3._r.append(fld_char_end)


def add_executive_summary(doc, domains, controls_by_domain):
    """Section 1: Executive Summary."""
    doc.add_heading("1. Executive Summary", level=1)

    # 1.1 Purpose
    doc.add_heading("1.1 Purpose of Assessment", level=2)
    doc.add_paragraph(
        "This independent assessment was conducted in accordance with SC-GL/6-2023 "
        "(Guidelines on Technology Risk Management) issued by the Securities Commission "
        "Malaysia. The purpose of this assessment is to evaluate the entity's technology "
        "risk management framework, policies, and controls across all domains prescribed "
        "by the guidelines, and to provide an independent opinion on the entity's "
        "compliance posture and control effectiveness."
    )
    doc.add_paragraph(
        "This report is intended for submission to the Securities Commission Malaysia "
        "as part of the entity's regulatory obligations and for the entity's board of "
        "directors and senior management to support governance and oversight of "
        "technology risk."
    )

    # 1.2 Scope
    doc.add_heading("1.2 Scope", level=2)
    doc.add_paragraph(
        "The assessment covered 35 controls across 10 technology risk management domains "
        "as defined in SC-GL/6-2023. The assessment was performed for the period "
        "[Start Date] to [End Date]."
    )

    p = doc.add_paragraph("Domains assessed:")
    for domain_slug in DOMAIN_ORDER:
        domain_info = domains[domain_slug]
        doc.add_paragraph(
            f"{domain_info['name']} ({domain_info['controlCount']} controls)",
            style="List Bullet",
        )

    # 1.3 Overall Assessment
    doc.add_heading("1.3 Overall Assessment", level=2)
    headers = ["Criterion", "Assessment"]
    rows = [
        ["Overall Compliance Score", "[X]%"],
        ["Overall Conclusion", "[Compliant / Partially Compliant / Non-Compliant]"],
        ["Key Strengths", "[Insert key strengths identified]"],
        ["Key Concerns", "[Insert key concerns identified]"],
    ]
    add_styled_table(doc, headers, rows)

    doc.add_paragraph("")

    # 1.4 Summary of Findings
    doc.add_heading("1.4 Summary of Findings", level=2)
    doc.add_paragraph(
        "The table below summarises the assessment results by domain. Detailed findings "
        "are presented in Section 3."
    )
    headers = [
        "Domain", "Controls", "Compliant", "Partial", "Non-Compliant", "N/A",
        "Domain Score",
    ]
    rows = []
    total_controls = 0
    for domain_slug in DOMAIN_ORDER:
        domain_info = domains[domain_slug]
        count = domain_info["controlCount"]
        total_controls += count
        rows.append([
            domain_info["name"],
            str(count),
            "[X]", "[X]", "[X]", "[X]",
            "[X]%",
        ])
    rows.append([
        "TOTAL", str(total_controls), "[X]", "[X]", "[X]", "[X]", "[X]%",
    ])
    add_styled_table(doc, headers, rows)


def add_methodology(doc):
    """Section 2: Methodology."""
    doc.add_heading("2. Methodology", level=1)

    # 2.1 Assessment Approach
    doc.add_heading("2.1 Assessment Approach", level=2)
    doc.add_paragraph(
        "This assessment was conducted in accordance with SC-GL/6-2023 Guidelines on "
        "Technology Risk Management and follows a structured assessment methodology "
        "designed to evaluate the design adequacy and operating effectiveness of "
        "technology risk management controls. The assessment utilised the Assessor Work "
        "Programme (AWP) developed specifically for SC-GL/6-2023 compliance reviews."
    )
    doc.add_paragraph(
        "The assessment approach comprised document review, stakeholder interviews, "
        "process walkthroughs, system demonstrations, and evidence inspection. "
        "Controls were assessed for both design adequacy (whether the control, if "
        "operating as designed, would mitigate the identified risk) and operating "
        "effectiveness (whether the control operated consistently over the assessment "
        "period)."
    )

    # 2.2 Assessment Methods
    doc.add_heading("2.2 Assessment Methods", level=2)
    headers = ["Method", "Description"]
    rows = [
        ["Inspection",
         "Examination of records, documents, policies, and configurations to "
         "verify existence, completeness, and compliance with requirements."],
        ["Inquiry",
         "Interviews and discussions with management, process owners, and "
         "operational staff to understand control design, implementation, "
         "and operation."],
        ["Observation",
         "Direct observation of processes, system demonstrations, and "
         "operational activities to verify controls operate as described."],
        ["Confirmation",
         "Independent verification of information through third-party "
         "confirmation, system testing, or corroboration from multiple "
         "sources."],
    ]
    add_styled_table(doc, headers, rows)

    doc.add_paragraph("")

    # 2.3 Conclusion Scale
    doc.add_heading("2.3 Conclusion Scale", level=2)
    doc.add_paragraph(
        "Each control was assessed and assigned one of the following conclusion ratings:"
    )
    headers = ["Rating", "Description"]
    rows = [
        ["Compliant",
         "The control is adequately designed and operating effectively. "
         "Evidence supports consistent compliance with SC-GL/6-2023 "
         "requirements throughout the assessment period."],
        ["Partially Compliant",
         "The control exists but has deficiencies in design or operation. "
         "Some requirements are met but gaps exist that require remediation "
         "to achieve full compliance."],
        ["Non-Compliant",
         "The control is absent, fundamentally deficient, or not operating. "
         "Significant gaps exist that present material risk to the entity's "
         "technology risk management posture."],
        ["N/A",
         "The control requirement is not applicable to the entity's "
         "operations (e.g., algorithmic trading controls for entities that "
         "do not engage in algorithmic trading). Rationale for N/A must be "
         "documented."],
    ]
    add_styled_table(doc, headers, rows)

    doc.add_paragraph("")

    # 2.4 Sampling Methodology
    doc.add_heading("2.4 Sampling Methodology", level=2)
    doc.add_paragraph(
        "Where assessment procedures required sample testing, samples were selected "
        "using a risk-based approach. For critical controls, a minimum sample of 5 "
        "instances was tested. For standard controls, a minimum sample of 3 instances "
        "was tested. Sampling methodology and sample sizes are documented in the "
        "Assessor Work Programme."
    )
    doc.add_paragraph(
        "Sample selection considered coverage across the assessment period, different "
        "transaction types, and both routine and non-routine events where applicable."
    )

    # 2.5 Limitations
    doc.add_heading("2.5 Limitations", level=2)
    doc.add_paragraph("The following limitations apply to this assessment:")
    limitations = [
        "This assessment provides limited assurance based on the evidence reviewed "
        "and procedures performed. It does not constitute an audit opinion.",
        "The assessment is based on information and evidence provided by the entity. "
        "The assessor has not independently verified the completeness or accuracy of "
        "all information provided.",
        "The assessment reflects the entity's compliance posture at a point in time "
        "and for the assessment period stated. Subsequent changes to controls, systems, "
        "or the threat landscape may affect the relevance of findings.",
        "Technical testing was limited to the scope defined in the engagement terms. "
        "The assessment does not represent a comprehensive penetration test or "
        "vulnerability assessment.",
        "The conclusions in this report are subject to the inherent limitations of "
        "any assessment process, including the possibility that material control "
        "weaknesses may not have been detected.",
    ]
    for i, limitation in enumerate(limitations, 1):
        doc.add_paragraph(f"{i}. {limitation}")


def add_detailed_findings(doc, domains, controls_by_domain, awp_lookup):
    """Section 3: Detailed Findings by Domain."""
    doc.add_heading("3. Detailed Findings by Domain", level=1)
    doc.add_paragraph(
        "This section presents the detailed assessment findings for each of the 10 "
        "technology risk management domains defined in SC-GL/6-2023. Each control is "
        "assessed individually with observations, evidence reviewed, recommendations, "
        "and management response."
    )

    for d_idx, domain_slug in enumerate(DOMAIN_ORDER, 1):
        domain_info = domains[domain_slug]
        domain_controls = controls_by_domain.get(domain_slug, [])

        # Domain heading
        doc.add_heading(f"3.{d_idx} {domain_info['name']}", level=2)
        doc.add_paragraph(domain_info["description"])

        for ctrl in domain_controls:
            slug = ctrl["slug"]
            awp = awp_lookup.get(slug, {})
            ref = awp.get("ref", "N/A")
            risk_tier = awp.get("riskTier", "standard").title()

            # Control heading
            doc.add_heading(f"{ref} — {ctrl['name']}", level=3)

            # Control summary table
            headers = ["Control Ref", "Risk Tier", "Conclusion", "Score"]
            rows = [
                [ref, risk_tier, "[Compliant / Partial / Non-Compliant / N/A]", "[X]%"],
            ]
            add_styled_table(doc, headers, rows)

            doc.add_paragraph("")

            # Control description
            p = doc.add_paragraph()
            run = p.add_run("Control Objective: ")
            run.font.bold = True
            run.font.color.rgb = DARK_BLUE
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run2 = p.add_run(ctrl["description"])
            run2.font.name = "Calibri"
            run2.font.size = Pt(10)
            run2.font.color.rgb = DARK_GREY

            # Observations
            p = doc.add_paragraph()
            run = p.add_run("Observations")
            run.font.bold = True
            run.font.color.rgb = DARK_BLUE
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            doc.add_paragraph("[Insert observations from assessment procedures performed. "
                              "Describe the current state of the control, evidence of "
                              "design adequacy, and operating effectiveness findings.]")

            # Evidence Reviewed
            p = doc.add_paragraph()
            run = p.add_run("Evidence Reviewed")
            run.font.bold = True
            run.font.color.rgb = DARK_BLUE
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            doc.add_paragraph("[List evidence documents reviewed during assessment. "
                              "Reference evidence IDs from Appendix A where applicable.]")

            # Recommendations
            p = doc.add_paragraph()
            run = p.add_run("Recommendations")
            run.font.bold = True
            run.font.color.rgb = DARK_BLUE
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            doc.add_paragraph("[Insert recommendations for improvement if applicable. "
                              "For compliant controls, state 'No recommendations — "
                              "control operating effectively.']")

            # Management Response
            p = doc.add_paragraph()
            run = p.add_run("Management Response")
            run.font.bold = True
            run.font.color.rgb = DARK_BLUE
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            doc.add_paragraph("[To be completed by entity management. Include agreement/"
                              "disagreement with finding, planned remediation actions, "
                              "responsible owner, and target completion date.]")


def add_maturity_assessment(doc, domains):
    """Section 4: Maturity Assessment."""
    doc.add_heading("4. Maturity Assessment", level=1)
    doc.add_paragraph(
        "This section presents the technology risk management maturity assessment "
        "across all domains. Maturity levels are assessed on a 0-3 scale as defined "
        "below."
    )

    # Maturity scale
    doc.add_heading("4.1 Maturity Scale", level=2)
    headers = ["Level", "Rating", "Description"]
    rows = [
        ["0", "Non-Existent",
         "No formal processes or controls exist. Technology risk management "
         "is ad-hoc and reactive."],
        ["1", "Initial",
         "Basic processes exist but are inconsistent and undocumented. "
         "Reliance on individual knowledge and effort."],
        ["2", "Defined",
         "Processes are documented, standardised, and consistently applied. "
         "Controls operate as designed with evidence of effectiveness."],
        ["3", "Advanced",
         "Processes are optimised, automated where appropriate, and "
         "continuously improved. Proactive risk management with measurable "
         "outcomes."],
    ]
    add_styled_table(doc, headers, rows)

    doc.add_paragraph("")

    # Domain maturity summary
    doc.add_heading("4.2 Domain Maturity Summary", level=2)
    headers = ["Domain", "Current Maturity", "Target Maturity", "Gap"]
    rows = []
    for domain_slug in DOMAIN_ORDER:
        domain_info = domains[domain_slug]
        rows.append([domain_info["name"], "[0-3]", "[0-3]", "[X]"])
    add_styled_table(doc, headers, rows)

    doc.add_paragraph("")

    # Maturity narrative
    doc.add_heading("4.3 Maturity Narrative", level=2)
    doc.add_paragraph(
        "[Provide a narrative assessment of the entity's overall technology risk "
        "management maturity. Highlight domains with the largest gaps between current "
        "and target maturity, and identify priority areas for improvement. Comment on "
        "the entity's maturity trajectory and readiness for evolving SC expectations.]"
    )


def add_recommendations_summary(doc):
    """Section 5: Recommendations Summary."""
    doc.add_heading("5. Recommendations Summary", level=1)
    doc.add_paragraph(
        "The table below consolidates all recommendations arising from this assessment, "
        "prioritised by severity. Detailed findings and context for each recommendation "
        "are provided in Section 3."
    )

    headers = ["#", "Finding", "Domain", "Severity", "Recommendation",
               "Target Date", "Status"]
    rows = [
        [str(i), "[Finding summary]", "[Domain]", "[Critical/High/Medium/Low]",
         "[Recommendation]", "[DD/MM/YYYY]", "[Open]"]
        for i in range(1, 6)
    ]
    add_styled_table(doc, headers, rows)


def add_management_response(doc):
    """Section 6: Management Response."""
    doc.add_heading("6. Management Response", level=1)
    doc.add_paragraph(
        "The entity's management has reviewed the findings and recommendations "
        "presented in this report. Management's response and commitment to "
        "remediation actions are documented below."
    )

    doc.add_heading("6.1 Management Statement", level=2)
    doc.add_paragraph(
        "[Insert management statement acknowledging the assessment findings and "
        "confirming commitment to addressing identified gaps. Management should "
        "confirm the accuracy of factual content and indicate any disagreements "
        "with findings or recommendations.]"
    )

    doc.add_paragraph("")

    doc.add_heading("6.2 Management Sign-Off", level=2)

    # Sign-off block
    headers = ["Field", "Details"]
    rows = [
        ["Name", "[Full Name]"],
        ["Title", "[Chief Executive Officer / Chief Technology Officer]"],
        ["Date", "[DD/MM/YYYY]"],
        ["Signature", ""],
    ]
    add_styled_table(doc, headers, rows)


def add_assessor_declaration(doc):
    """Section 7: Assessor Declaration."""
    doc.add_heading("7. Assessor Declaration", level=1)

    doc.add_paragraph(
        "We hereby declare that this independent assessment of the entity's compliance "
        "with SC-GL/6-2023 Guidelines on Technology Risk Management has been conducted "
        "in accordance with the following principles:"
    )

    declarations = [
        "Independence: The assessment team maintained independence from the entity "
        "throughout the engagement. No member of the assessment team has a financial "
        "interest in, or other relationship with, the entity that could impair "
        "objectivity.",
        "Professional Standards: The assessment was conducted in accordance with "
        "applicable professional standards and the engagement terms agreed with "
        "the entity.",
        "Limited Assurance: This assessment provides limited assurance based on "
        "the procedures performed and evidence reviewed. It does not constitute "
        "an audit opinion or absolute assurance on the effectiveness of the "
        "entity's technology risk management framework.",
        "Competence: The assessment team possesses the necessary qualifications, "
        "experience, and competence in technology risk management and capital "
        "market operations to conduct this assessment.",
        "Confidentiality: All information obtained during this assessment has been "
        "treated as confidential and will not be disclosed to third parties without "
        "the entity's consent, except as required by law or regulatory obligation.",
    ]
    for decl in declarations:
        doc.add_paragraph(decl, style="List Bullet")

    doc.add_paragraph("")

    # Lead Assessor sign-off
    doc.add_heading("Lead Assessor", level=2)
    headers = ["Field", "Details"]
    rows = [
        ["Name", "[Full Name]"],
        ["Designation", "[Title / Position]"],
        ["Firm", "[Assessor Firm Name]"],
        ["Professional Qualifications", "[e.g., CISSP, CISA, CISM]"],
        ["Date", "[DD/MM/YYYY]"],
        ["Signature", ""],
    ]
    add_styled_table(doc, headers, rows)

    doc.add_paragraph("")

    # Quality Reviewer sign-off
    doc.add_heading("Quality Reviewer", level=2)
    headers = ["Field", "Details"]
    rows = [
        ["Name", "[Full Name]"],
        ["Designation", "[Title / Position]"],
        ["Firm", "[Assessor Firm Name]"],
        ["Date", "[DD/MM/YYYY]"],
        ["Signature", ""],
    ]
    add_styled_table(doc, headers, rows)


def add_appendix_evidence(doc):
    """Appendix A: Evidence Index."""
    doc.add_heading("Appendix A: Evidence Index", level=1)
    doc.add_paragraph(
        "The following table indexes all evidence documents reviewed during this "
        "assessment. Evidence is maintained in the assessor's working papers file."
    )

    headers = ["#", "Evidence ID", "Document", "Source", "Date Received", "Status"]
    rows = [
        [str(i), f"EV-{i:03d}", "[Document name]", "[Source]", "[DD/MM/YYYY]",
         "[Received / Pending / N/A]"]
        for i in range(1, 11)
    ]
    add_styled_table(doc, headers, rows)


def add_appendix_glossary(doc):
    """Appendix B: Glossary."""
    doc.add_heading("Appendix B: Glossary", level=1)
    doc.add_paragraph(
        "The following abbreviations and terms are used throughout this report."
    )

    headers = ["Abbreviation", "Full Term"]
    glossary = [
        ("AI", "Artificial Intelligence"),
        ("AWP", "Assessor Work Programme"),
        ("BCP", "Business Continuity Plan"),
        ("BNM", "Bank Negara Malaysia"),
        ("Bursa", "Bursa Malaysia"),
        ("CISO", "Chief Information Security Officer"),
        ("CSP", "Cloud Service Provider"),
        ("CTO", "Chief Technology Officer"),
        ("DLP", "Data Loss Prevention"),
        ("DLT", "Distributed Ledger Technology"),
        ("DR", "Disaster Recovery"),
        ("GTRM", "Guidelines on Technology Risk Management"),
        ("IDS/IPS", "Intrusion Detection System / Intrusion Prevention System"),
        ("KRI", "Key Risk Indicator"),
        ("ML", "Machine Learning"),
        ("PDPA", "Personal Data Protection Act 2010"),
        ("RBAC", "Role-Based Access Control"),
        ("RPO", "Recovery Point Objective"),
        ("RTO", "Recovery Time Objective"),
        ("SC", "Securities Commission Malaysia"),
        ("SC-GL/6-2023", "SC Guidelines on Technology Risk Management (2023)"),
        ("SIEM", "Security Information and Event Management"),
        ("SLA", "Service Level Agreement"),
        ("SOAR", "Security Orchestration, Automation and Response"),
        ("SOC", "Security Operations Centre"),
    ]
    rows = [[abbr, full] for abbr, full in glossary]
    add_styled_table(doc, headers, rows)


def main():
    """Build the SC-GTRM assessment report template."""
    print("Loading data...")
    library, domains, awp_lookup, controls_by_domain = load_data()

    print("Creating document...")
    doc = Document()
    configure_styles(doc)
    set_margins(doc)

    # Cover Page
    add_cover_page(doc)
    add_page_break(doc)

    # Document Control
    add_document_control(doc)
    add_page_break(doc)

    # Table of Contents
    add_toc(doc)
    add_page_break(doc)

    # Section 1: Executive Summary
    add_executive_summary(doc, domains, controls_by_domain)
    add_page_break(doc)

    # Section 2: Methodology
    add_methodology(doc)
    add_page_break(doc)

    # Section 3: Detailed Findings
    add_detailed_findings(doc, domains, controls_by_domain, awp_lookup)
    add_page_break(doc)

    # Section 4: Maturity Assessment
    add_maturity_assessment(doc, domains)
    add_page_break(doc)

    # Section 5: Recommendations Summary
    add_recommendations_summary(doc)
    add_page_break(doc)

    # Section 6: Management Response
    add_management_response(doc)
    add_page_break(doc)

    # Section 7: Assessor Declaration
    add_assessor_declaration(doc)
    add_page_break(doc)

    # Appendix A: Evidence Index
    add_appendix_evidence(doc)
    add_page_break(doc)

    # Appendix B: Glossary
    add_appendix_glossary(doc)

    # Headers and Footers (applied after all content)
    add_header_footer(doc)

    # Save
    print(f"Saving to {OUTPUT}...")
    doc.save(str(OUTPUT))
    print(f"Report template generated: {OUTPUT}")
    print(f"  - 10 domains, 35 controls")
    print(f"  - {len(awp_lookup)} controls with AWP risk tier data")


if __name__ == "__main__":
    main()
